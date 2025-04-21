# c:\Users\user\Documents\gen-game\Agents\Tutor.py
import os
import re
import sys
import json
import time
import asyncio
import threading
import traceback
import tkinter as tk
from tkinter import filedialog, messagebox
import google.generativeai as genai
from typing import List, Dict, Any
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler, FileSystemEvent
import mammoth
import datetime,textwrap
from collections import defaultdict
import hashlib # Added for hashing
sys.path.append(os.path.abspath(os.path.dirname(__file__)))
from watchdog_tool import comprehensive_scene_analysis
import msvcrt

GREEN = "\033[92m"
BLUE = "\033[94m"
CYAN = "\033[96m"
RED = "\033[91m"
RESET = "\033[0m"
YELLOW = "\033[93m"
MAGENTA = "\033[95m"

VALIDATION_RESULTS_FILE = 'gdd_validation_results.json'
GDD_HASH_FILE = 'gdd_hash.txt' # Added for storing GDD hash

try:
    import docx
except ImportError:
    print(f"{YELLOW}python-docx not installed. Will use text file for reports instead.{RESET}")
    print(f"{BLUE}To install: pip install python-docx{RESET}")

# --- GDD Validation and Report Functions ---
def validate_gdd_with_octalysis(gemini_client, gdd_content: str) -> Dict:
    try:
        prompt = f"""
        You are a game design expert specializing in gamification through the Octalysis Framework.

        Review and enhance the following Game Design Document for a dyslexia support educational game:

        {gdd_content}

        Instructions:
        1. COMPLETELY fill ALL sections using Octalysis Framework - leave nothing empty
        2. For any undefined original sections (like core mechanics), create new content
        3. Ensure ALL 8 core drives are explicitly implemented:
           - Epic Meaning & Calling
           - Development & Accomplishment
           - Empowerment of Creativity & Feedback
           - Ownership & Possession
           - Social Influence & Relatedness
           - Scarcity & Impatience
           - Unpredictability & Curiosity
           - Loss & Avoidance
        4. Address ALL player types with specific features
        5. Include concrete examples for each section
        6. Maintain original structure but enhance all components

        Format STRICTLY follow this template:

        ------
        OVERALL_SCORE: [X.X/10]

        TOP_RECOMMENDATIONS:
        - Recommendation 1
        - Recommendation 2
        - Recommendation 3

        SUMMARY: [2-3 sentence summary]

        TITLE: [Enhanced Title]

        ORIGINAL_CONCEPT: [Original concept summary]

        OVERVIEW: [Enhanced overview paragraph]

        CORE_MECHANICS:
        - [Mechanic 1] (ties to [Core Drive])
        - [Mechanic 2] (ties to [Core Drive])
        - [Mechanic 3] (ties to [Core Drive])

        PROGRESSION_SYSTEM:
        [Detailed progression description with core drives]

        REWARD_SYSTEMS:
        - [Reward 1] (motivates [Core Drive])
        - [Reward 2] (motivates [Core Drive])
        - [Reward 3] (motivates [Core Drive])

        SOCIAL_ELEMENTS:
        - [Feature 1] (supports [Core Drive])
        - [Feature 2] (supports [Core Drive])
        - [Feature 3] (supports [Core Drive])

        PLAYER_JOURNEY:

        DISCOVERY:
        - [Feature 1]
        - [Feature 2]
        - [Feature 3]

        ONBOARDING:
        - [Feature 1]
        - [Feature 2]
        - [Feature 3]

        SCAFFOLDING:
        - [Feature 1]
        - [Feature 2]
        - [Feature 3]

        ENDGAME:
        - [Feature 1]
        - [Feature 2]
        - [Feature 3]

        PLAYER_TYPES:

        ACHIEVERS:
        - [Feature 1] (supports [Core Drive])
        - [Feature 2] (supports [Core Drive])

        EXPLORERS:
        - [Feature 1] (supports [Core Drive])
        - [Feature 2] (supports [Core Drive])

        SOCIALIZERS:
        - [Feature 1] (supports [Core Drive])
        - [Feature 2] (supports [Core Drive])

        COMPETITORS:
        - [Feature 1] (supports [Core Drive])
        - [Feature 2] (supports [Core Drive])

        DYSLEXIA_SUPPORT_FEATURES:
        - [Feature 1] (supports [Need])
        - [Feature 2] (supports [Need])
        - [Feature 3] (supports [Need])

        IMPLEMENTATION_ROADMAP:
        1. [Phase 1]
        2. [Phase 2]
        3. [Phase 3]
        4. [Phase 4]
        5. [Phase 5]
        ------
        """

        response = gemini_client.generate_content(prompt)

        def extract_section(pattern, text, default):
            match = re.search(pattern, text, re.DOTALL)
            return match.group(1).strip() if match else default

        def extract_list(pattern, text):
            match = re.search(pattern, text, re.DOTALL)
            return [line.strip() for line in match.group(1).split('\n') if line.strip()] if match else []

        content = response.text
        sections = {
            "overall_score": float(extract_section(r'OVERALL_SCORE:\s*(\d+\.?\d*)', content, '0.0')),
            "top_recommendations": extract_list(r'TOP_RECOMMENDATIONS:(.*?)(?=\n[A-Z]+:)', content),
            "summary": extract_section(r'SUMMARY:(.*?)(?=\n[A-Z]+:)', content, "No summary provided"),
            "title": extract_section(r'TITLE:(.*?)(?=\n[A-Z]+:)', content, "Enhanced Game Design Document"),
            "original_concept": extract_section(r'ORIGINAL_CONCEPT:(.*?)(?=\n[A-Z]+:)', content, "Original concept not specified"),
            "enhanced_gdd": {
                "overview": extract_section(r'OVERVIEW:(.*?)(?=\n[A-Z]+:)', content, "No overview provided"),
                "core_mechanics": extract_list(r'CORE_MECHANICS:(.*?)(?=\n[A-Z]+:)', content),
                "progression_system": extract_section(r'PROGRESSION_SYSTEM:(.*?)(?=\n[A-Z]+:)', content, "Progression system not defined"),
                "reward_systems": extract_list(r'REWARD_SYSTEMS:(.*?)(?=\n[A-Z]+:)', content),
                "social_elements": extract_list(r'SOCIAL_ELEMENTS:(.*?)(?=\n[A-Z]+:)', content),
                "player_journey": {
                    "discovery": extract_list(r'DISCOVERY:(.*?)(?=\n[A-Z]+:)', content),
                    "onboarding": extract_list(r'ONBOARDING:(.*?)(?=\n[A-Z]+:)', content),
                    "scaffolding": extract_list(r'SCAFFOLDING:(.*?)(?=\n[A-Z]+:)', content),
                    "endgame": extract_list(r'ENDGAME:(.*?)(?=\n[A-Z]+:)', content)
                },
                "player_types": {
                    "achievers": extract_list(r'ACHIEVERS:(.*?)(?=\n[A-Z]+:)', content),
                    "explorers": extract_list(r'EXPLORERS:(.*?)(?=\n[A-Z]+:)', content),
                    "socializers": extract_list(r'SOCIALIZERS:(.*?)(?=\n[A-Z]+:)', content),
                    "competitors": extract_list(r'COMPETITORS:(.*?)(?=\n[A-Z]+:)', content)
                },
                "dyslexia_support_features": extract_list(r'DYSLEXIA_SUPPORT_FEATURES:(.*?)(?=\n[A-Z]+:)', content),
                "implementation_roadmap": extract_list(r'IMPLEMENTATION_ROADMAP:(.*?)(?=\n[A-Z]+:|\Z)', content)
            }
        }

        for section in ['core_mechanics', 'reward_systems', 'social_elements']:
            if not sections['enhanced_gdd'][section]:
                sections['enhanced_gdd'][section] = [f"Automatically generated {section.replace('_', ' ')} content"]

        return sections

    except Exception as e:
        return {"error": str(e)}

def save_validation_report(validation_results: Dict, output_path: str) -> str:
    report_doc_path = None
    results_json_path = os.path.join(output_path, VALIDATION_RESULTS_FILE)

    try:
        with open(results_json_path, 'w') as f:
            json.dump(validation_results, f, indent=4)
        print(f"{GREEN}Validation results saved to: {results_json_path}{RESET}")
    except Exception as e:
        print(f"{RED}Error saving validation results JSON: {e}{RESET}")

    try:
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = docx.Document()

        def add_heading(text, level):
            heading = doc.add_heading(text, level)
            if heading.style.font:
                 heading.style.font.size = Pt(14 if level == 1 else 12)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def add_section(title, content, level=1):
            add_heading(title, level)
            if isinstance(content, list):
                if not content:
                     doc.add_paragraph("N/A", style='BodyText')
                else:
                    for item in content:
                        doc.add_paragraph(f"• {str(item)}", style='ListBullet')
            elif content:
                doc.add_paragraph(str(content))
            else:
                doc.add_paragraph("N/A", style='BodyText')

        add_heading(validation_results.get('title', 'Enhanced GDD'), 0)
        add_section('Octalysis Score', f"{validation_results.get('overall_score', 'N/A')}/10")
        add_section('Top Recommendations', validation_results.get('top_recommendations', []))
        add_section('Executive Summary', validation_results.get('summary', ''))
        add_section('Original Concept', validation_results.get('original_concept', ''))

        enhanced = validation_results.get('enhanced_gdd', {})
        add_section('Enhanced Overview', enhanced.get('overview', ''))
        add_section('Core Mechanics', enhanced.get('core_mechanics', []))
        add_section('Progression System', enhanced.get('progression_system', ''))
        add_section('Reward Systems', enhanced.get('reward_systems', []))
        add_section('Social Elements', enhanced.get('social_elements', []))

        add_heading('Player Journey', 1)
        journey = enhanced.get('player_journey', {})
        for phase in ['Discovery', 'Onboarding', 'Scaffolding', 'Endgame']:
            add_section(phase, journey.get(phase.lower(), []), 2)

        add_heading('Player Types', 1)
        types = enhanced.get('player_types', {})
        for ptype in ['Achievers', 'Explorers', 'Socializers', 'Competitors']:
            add_section(ptype, types.get(ptype.lower(), []), 2)

        add_section('Dyslexia Support', enhanced.get('dyslexia_support_features', []))
        add_section('Implementation Roadmap', enhanced.get('implementation_roadmap', []))

        report_doc_path = os.path.join(output_path, 'Enhanced_GDD_Full.docx')
        doc.save(report_doc_path)
        print(f"{GREEN}Enhanced GDD report saved to: {report_doc_path}{RESET}")
        return report_doc_path

    except ImportError:
        print(f"{YELLOW}python-docx not found. Skipping Word document generation.{RESET}")
        return results_json_path if os.path.exists(results_json_path) else None
    except Exception as e:
        print(f"{RED}Error saving Word document: {e}{RESET}")
        return results_json_path if os.path.exists(results_json_path) else None

# --- File Selection and GDD Extraction ---
def select_project_path() -> str:
    root = tk.Tk()
    root.withdraw()
    project_path = filedialog.askdirectory(
        title="Select Unity Project Directory",
        initialdir=os.path.expanduser("~")
    )
    return project_path

def select_gdd_file() -> str:
    root = tk.Tk()
    root.withdraw()
    gdd_file = filedialog.askopenfilename(
        title="Select Game Design Document",
        initialdir=os.path.expanduser("~"),
        filetypes=[
            ("Word Documents", "*.docx"),
            ("PDF Files", "*.pdf"),
            ("Text Files", "*.txt"),
            ("All Files", "*.*")
        ]
    )
    return gdd_file

def extract_gdd_content(gdd_path: str) -> str:
    try:
        if gdd_path.lower().endswith('.docx'):
            with open(gdd_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
        elif gdd_path.lower().endswith('.txt'):
            with open(gdd_path, 'r', encoding='utf-8') as txt_file:
                return txt_file.read()
        elif gdd_path.lower().endswith('.pdf'):
            try:
                import PyPDF2
                with open(gdd_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    return ' '.join(page.extract_text() for page in pdf_reader.pages)
            except ImportError:
                print(f"{RED}PyPDF2 library not installed. Cannot read PDF.{RESET}")
                return ""
        else:
            print(f"{RED}Unsupported file format: {gdd_path}{RESET}")
            return ""
    except Exception as e:
        print(f"{RED}Error extracting GDD content: {e}{RESET}")
        print(f"{BLUE}Traceback:{RESET}")
        import traceback
        traceback.print_exc()
        return ""

# --- GDD Hashing Function ---
def calculate_gdd_hash(gdd_content: str) -> str:
    """Calculates the SHA256 hash of the GDD content."""
    return hashlib.sha256(gdd_content.encode('utf-8')).hexdigest()

# --- Metrics Logger Class ---
class MetricsLogger:
    def __init__(self, project_path: str):
        self.metrics_file = os.path.join(project_path, 'learning_metrics.md')
        self.metrics = {
            'start_time': datetime.datetime.now().isoformat(),
            'end_time': None,
            'gdd_refinements': 0,
            'tasks': defaultdict(dict),
            'chapters': defaultdict(dict),
            'coin_transactions': [],
            'script_purchases': 0,
            'scene_analysis_count': 0,
            'validation_attempts': defaultdict(int),
            'menu_accesses': 0,
            'current_coins': 0,
            'errors': [],
            'events': [],
            'learning_path': {} # Initialize learning_path key here
        }
        # Initial save might not have learning path yet, which is fine
        self._save_metrics()

    def log_task_start(self, chapter_num: int, task_num: int, task_title: str):
        task_key = f"chapter_{chapter_num}_task_{task_num}"
        self.metrics['tasks'][task_key] = {
            'title': task_title,
            'start_time': datetime.datetime.now().isoformat(),
            'end_time': None,
            'completed': False,
            'validation_attempts': 0
        }
        self._save_metrics()

    def log_task_completion(self, chapter_num: int, task_num: int, success: bool):
        task_key = f"chapter_{chapter_num}_task_{task_num}"
        if task_key in self.metrics['tasks']:
            self.metrics['tasks'][task_key]['end_time'] = datetime.datetime.now().isoformat()
            self.metrics['tasks'][task_key]['completed'] = success
            if 'start_time' in self.metrics['tasks'][task_key] and self.metrics['tasks'][task_key]['start_time']:
                try:
                    start = datetime.datetime.fromisoformat(self.metrics['tasks'][task_key]['start_time'])
                    end = datetime.datetime.now()
                    self.metrics['tasks'][task_key]['duration_seconds'] = (end - start).total_seconds()
                except ValueError:
                     self.metrics['tasks'][task_key]['duration_seconds'] = None # Handle invalid start time format
        self._save_metrics()

    def log_chapter_start(self, chapter_num: int, chapter_title: str):
        chapter_key = f"chapter_{chapter_num}"
        self.metrics['chapters'][chapter_key] = {
            'title': chapter_title,
            'start_time': datetime.datetime.now().isoformat(),
            'end_time': None,
            'completed': False
        }
        self._save_metrics()

    def log_chapter_completion(self, chapter_num: int):
        chapter_key = f"chapter_{chapter_num}"
        if chapter_key in self.metrics['chapters']:
            self.metrics['chapters'][chapter_key]['end_time'] = datetime.datetime.now().isoformat()
            self.metrics['chapters'][chapter_key]['completed'] = True
            if 'start_time' in self.metrics['chapters'][chapter_key] and self.metrics['chapters'][chapter_key]['start_time']:
                try:
                    start = datetime.datetime.fromisoformat(self.metrics['chapters'][chapter_key]['start_time'])
                    end = datetime.datetime.now()
                    self.metrics['chapters'][chapter_key]['duration_seconds'] = (end - start).total_seconds()
                except ValueError:
                    self.metrics['chapters'][chapter_key]['duration_seconds'] = None # Handle invalid start time format
        self._save_metrics()

    def log_gdd_refinement(self):
        self.metrics['gdd_refinements'] += 1
        self._save_metrics()

    def log_coin_transaction(self, amount: int, reason: str, current_balance: int):
        self.metrics['coin_transactions'].append({
            'timestamp': datetime.datetime.now().isoformat(),
            'amount': amount,
            'reason': reason,
            'balance': current_balance
        })
        self.metrics['current_coins'] = current_balance
        self._save_metrics()

    def log_script_purchase(self):
        self.metrics['script_purchases'] += 1
        self._save_metrics()

    def log_scene_analysis(self):
        self.metrics['scene_analysis_count'] += 1
        self._save_metrics()

    def log_validation_attempt(self, chapter_num: int, task_num: int):
        task_key = f"chapter_{chapter_num}_task_{task_num}"
        self.metrics['validation_attempts'][task_key] += 1
        if task_key in self.metrics['tasks']:
            self.metrics['tasks'][task_key]['validation_attempts'] = self.metrics['validation_attempts'][task_key]
        self._save_metrics()

    def log_menu_access(self):
        self.metrics['menu_accesses'] += 1
        self._save_metrics()

    def log_error(self, error: str):
        self.metrics['errors'].append({
            'timestamp': datetime.datetime.now().isoformat(),
            'error': error
        })
        self._save_metrics()

    def log_session_end(self):
        self.metrics['end_time'] = datetime.datetime.now().isoformat()
        if 'start_time' in self.metrics and self.metrics['start_time']:
            try:
                start = datetime.datetime.fromisoformat(self.metrics['start_time'])
                end = datetime.datetime.now()
                self.metrics['total_duration_seconds'] = (end - start).total_seconds()
            except ValueError:
                 self.metrics['total_duration_seconds'] = None # Handle invalid start time format
        self._save_metrics()

    def log_event(self, event_name: str, details: Dict = None):
        event_data = {
            'timestamp': datetime.datetime.now().isoformat(),
            'event': event_name
        }
        if details:
            event_data['details'] = details
        self.metrics['events'].append(event_data)
        self._save_metrics()

    def _save_metrics(self):
        """Saves the current metrics dictionary to the markdown file."""
        # This method now just needs to format whatever is in self.metrics
        try:
            with open(self.metrics_file, 'w') as f:
                f.write(self._format_metrics())
        except Exception as e:
            # Avoid infinite loops if logging the error causes another save attempt
            print(f"{RED}Critical Error saving metrics: {e}{RESET}")
            # Optionally log to stderr or a separate critical log file
            # import sys
            # print(f"Critical Error saving metrics: {e}", file=sys.stderr)

    # --- CORRECTED _format_metrics ---
    def _format_metrics(self) -> str:
        """Formats the metrics dictionary into a markdown string."""
        md = "# Unity Learning Tutor Metrics Report\n\n"
        md += f"## Session Overview\n"
        md += f"- **Start Time**: {self.metrics.get('start_time', 'N/A')}\n"
        md += f"- **End Time**: {self.metrics.get('end_time', 'Session in progress')}\n"
        if 'total_duration_seconds' in self.metrics and self.metrics['total_duration_seconds'] is not None:
            md += f"- **Total Duration**: {self.metrics['total_duration_seconds']:.2f} seconds\n"
        md += f"- **GDD Refinements**: {self.metrics.get('gdd_refinements', 0)}\n"
        md += f"- **Script Purchases**: {self.metrics.get('script_purchases', 0)}\n"
        md += f"- **Scene Analyses**: {self.metrics.get('scene_analysis_count', 0)}\n"
        md += f"- **Menu Accesses**: {self.metrics.get('menu_accesses', 0)}\n"
        md += f"- **Current Coin Balance**: {self.metrics.get('current_coins', 0)}\n\n"

        md += "## Events\n"
        events = self.metrics.get('events', [])
        if events:
            md += "| Timestamp | Event | Details |\n"
            md += "|-----------|-------|---------|\n"
            for event in events:
                details_str = json.dumps(event.get('details', {}))
                md += f"| {event.get('timestamp', 'N/A')} | {event.get('event', 'N/A')} | {details_str} |\n"
        else:
            md += "No events recorded.\n"
        md += "\n"

        md += "## Coin Transactions\n"
        transactions = self.metrics.get('coin_transactions', [])
        if transactions:
            md += "| Timestamp | Amount | Reason | Balance |\n"
            md += "|-----------|--------|--------|---------|\n"
            for tx in transactions:
                md += f"| {tx.get('timestamp', 'N/A')} | {tx.get('amount', 'N/A')} | {tx.get('reason', 'N/A')} | {tx.get('balance', 'N/A')} |\n"
        else:
            md += "No coin transactions recorded.\n"
        md += "\n"

        md += "## Chapter Progress\n"
        chapters = self.metrics.get('chapters', {})
        if chapters:
            for chapter_key, chapter_data in chapters.items():
                md += f"### {chapter_key.replace('_', ' ').title()}: {chapter_data.get('title', '')}\n"
                md += f"- Started: {chapter_data.get('start_time', 'N/A')}\n"
                md += f"- Completed: {chapter_data.get('end_time', 'Not completed')}\n"
                if 'duration_seconds' in chapter_data and chapter_data['duration_seconds'] is not None:
                    md += f"- Duration: {chapter_data['duration_seconds']:.2f} seconds\n"
                md += "\n"
        else:
            md += "No chapter progress recorded.\n"
        md += "\n"

        md += "## Task Details\n"
        tasks = self.metrics.get('tasks', {})
        if tasks:
            for task_key, task_data in tasks.items():
                md += f"### {task_key.replace('_', ' ').title()}: {task_data.get('title', '')}\n"
                md += f"- Started: {task_data.get('start_time', 'N/A')}\n"
                md += f"- Completed: {task_data.get('end_time', 'Not completed')}\n"
                md += f"- Success: {task_data.get('completed', 'N/A')}\n"
                md += f"- Validation Attempts: {task_data.get('validation_attempts', 0)}\n"
                if 'duration_seconds' in task_data and task_data['duration_seconds'] is not None:
                    md += f"- Duration: {task_data['duration_seconds']:.2f} seconds\n"
                md += "\n"
        else:
            md += "No task details recorded.\n"
        md += "\n"

        md += "## Errors\n"
        errors = self.metrics.get('errors', [])
        if errors:
            md += "| Timestamp | Error |\n"
            md += "|-----------|-------|\n"
            for error in errors:
                md += f"| {error.get('timestamp', 'N/A')} | {error.get('error', 'N/A')} |\n"
        else:
            md += "No errors recorded.\n"
        md += "\n"

        # --- CORRECTED SECTION ---
        md += "## Learning Path\n"
        md += "```json\n"
        # Get learning_path from the metrics dictionary using .get()
        learning_path_data = self.metrics.get('learning_path', {}) # Default to empty dict
        if learning_path_data: # Check if the retrieved data is not empty/None
            try:
                # Dump the data retrieved from the dictionary
                md += json.dumps(learning_path_data, indent=2)
            except TypeError as e:
                md += f"Error formatting learning path: {e}"
        else:
            # Output empty JSON if no learning path in metrics or it's empty
            md += "{}"
        md += "\n```\n"
        # --- END OF CORRECTED SECTION ---

        return md
    # --- END OF CORRECTED _format_metrics ---

# --- Unity Learning Tutor Class ---
class UnityLearningTutor:
    # Modified __init__ to accept validation results and flag
    def __init__(self, gemini_api_key: str, project_path: str, gdd_content: str,
                 validation_results: Dict | None, gdd_was_validated: bool):
        genai.configure(api_key=gemini_api_key)
        self.gemini_client = genai.GenerativeModel('gemini-2.5-flash-preview-04-17')
        self.project_path = project_path
        self.metrics_logger = MetricsLogger(project_path) # Init logger early
        self.learning_path = {}
        self.current_chapter_index = 0
        self.current_task_index = 0
        self.gdd_content = gdd_content # Store the initial GDD content
        self.user_coins = 0
        self.script_detail_level = "Give only the functions and description of each one"
        self.coin_rewards = {
            "task_completion": 5,
            "chapter_completion": 15
        }
        self.coin_costs = {
            "full_script": 10
        }
        self.progress_file = os.path.join(project_path, 'dyslexia_game_progress.json')
        self.gdd_hash_file_path = os.path.join(project_path, GDD_HASH_FILE) # Path for hash file
        self._scene_change_flag = threading.Event()
        self.watchdog = None
        self.watchdog_observer = None
        self.scene_analysis_count = 0

        # Store results and flag from main()
        self.validation_results = validation_results
        self.gdd_was_validated = gdd_was_validated # True if GDD changed and was re-validated in main()

        # Load progress attempts to set initial state if available
        self.load_progress() # This will now update metrics_logger.metrics['learning_path']

    def save_gdd_hash(self, gdd_content: str):
        """Saves the hash of the current GDD content."""
        try:
            current_hash = calculate_gdd_hash(gdd_content)
            with open(self.gdd_hash_file_path, 'w') as f:
                f.write(current_hash)
            print(f"{BLUE}Saved current GDD hash.{RESET}")
        except Exception as e:
            print(f"{RED}Error saving GDD hash: {e}{RESET}")
            self.metrics_logger.log_error(f"GDD Hash Save Error: {e}")

    def load_gdd_hash(self) -> str | None:
        """Loads the previously saved GDD hash."""
        if os.path.exists(self.gdd_hash_file_path):
            try:
                with open(self.gdd_hash_file_path, 'r') as f:
                    return f.read().strip()
            except Exception as e:
                print(f"{RED}Error loading GDD hash: {e}{RESET}")
                self.metrics_logger.log_error(f"GDD Hash Load Error: {e}")
                return None
        return None # No hash file found

    def load_gdd_validation_results(self) -> Dict | None:
        """Loads GDD validation results from the project directory."""
        results_path = os.path.join(self.project_path, VALIDATION_RESULTS_FILE)
        if os.path.exists(results_path):
            try:
                with open(results_path, 'r') as f:
                    return json.load(f)
            except Exception as e:
                print(f"{RED}Error loading GDD validation results: {e}{RESET}")
                self.metrics_logger.log_error(f"GDD Results Load Error: {e}")
                return None
        else:
            return None

    async def enhance_learning_path_with_gamification(self, validation_results: Dict):
        if not self.learning_path or 'chapters' not in self.learning_path:
            print(f"{YELLOW}Cannot enhance: Learning path is empty or invalid.{RESET}")
            return
        if not validation_results or 'enhanced_gdd' not in validation_results:
            print(f"{YELLOW}Cannot enhance: GDD validation results are missing or invalid.{RESET}")
            return

        print(f"\n{CYAN}Enhancing remaining learning path with GDD gamification...{RESET}")
        self.metrics_logger.log_event("gamification_enhancement_start")

        enhanced_gdd = validation_results['enhanced_gdd']
        gamification_elements = {
            "Core Mechanics": enhanced_gdd.get('core_mechanics', []),
            "Reward Systems": enhanced_gdd.get('reward_systems', []),
            "Social Elements": enhanced_gdd.get('social_elements', []),
            "Player Type Features (Achievers)": enhanced_gdd.get('player_types', {}).get('achievers', []),
            "Player Type Features (Explorers)": enhanced_gdd.get('player_types', {}).get('explorers', []),
            "Player Type Features (Socializers)": enhanced_gdd.get('player_types', {}).get('socializers', []),
            "Player Type Features (Competitors)": enhanced_gdd.get('player_types', {}).get('competitors', []),
        }

        all_recommendations = []
        for category, items in gamification_elements.items():
            if isinstance(items, list):
                all_recommendations.extend([f"{category}: {item}" for item in items])
            elif isinstance(items, str) and items:
                 all_recommendations.append(f"{category}: {items}")

        if not all_recommendations:
             print(f"{YELLOW}No specific gamification elements found in GDD results to inject.{RESET}")
             self.metrics_logger.log_event("gamification_enhancement_skipped", details={"reason": "no elements found"})
             return

        start_chap_idx = self.current_chapter_index
        start_task_idx = self.current_task_index

        tasks_to_enhance = []
        for i, chapter in enumerate(self.learning_path['chapters']):
            if i < start_chap_idx:
                continue
            for j, task in enumerate(chapter['tasks']):
                if i == start_chap_idx and j < start_task_idx:
                    continue
                tasks_to_enhance.append({'chapter_index': i, 'task_index': j, 'task_data': task})

        if not tasks_to_enhance:
            print(f"{YELLOW}No remaining tasks to enhance.{RESET}")
            self.metrics_logger.log_event("gamification_enhancement_skipped", details={"reason": "no remaining tasks"})
            return

        print(f"{BLUE}Found {len(tasks_to_enhance)} remaining tasks to potentially enhance.{RESET}")

        updated_tasks_count = 0
        api_call_delay = 1.1

        for task_info in tasks_to_enhance:
            chap_idx = task_info['chapter_index']
            task_idx = task_info['task_index']
            original_task = task_info['task_data']

            prompt = f"""
            You are a Unity Learning Path Gamification Expert.
            Your goal is to subtly enhance an existing learning task to incorporate relevant gamification elements based on a GDD analysis (using Octalysis).

            GDD Gamification Recommendations (Octalysis-based):
            ---
            {chr(10).join(all_recommendations)}
            ---

            Original Task Details:
            ---
            Chapter: {self.learning_path['chapters'][chap_idx]['title']} (Number: {self.learning_path['chapters'][chap_idx]['number']})
            Task Number: {original_task['number']}
            Description: {original_task['description']}
            Explanation: {original_task.get('explanation', 'N/A')}
            Steps:
            {chr(10).join([f'- {s}' for s in original_task.get('steps', [])])}
            Expected Object: {json.dumps(original_task.get('expected_object', {}))}
            ---

            Instructions:
            1. Review the Original Task Details and the GDD Gamification Recommendations.
            2. Identify 1-2 *relevant* gamification elements from the recommendations that could be naturally integrated into THIS specific task. Focus on relevance (e.g., add reward feedback to a task completion, add exploration elements to a scene setup task).
            3. Modify the task's 'description', 'explanation', or 'steps' to incorporate these chosen elements. Changes should be subtle and enhance the existing task, not create entirely new unrelated sub-tasks.
            4. Examples of subtle changes:
                - Add a step: "Add a simple particle effect and sound when the letter is correctly matched (Reward System - Development & Accomplishment)."
                - Modify explanation: "Setting up the environment encourages exploration (Player Type - Explorers)."
                - Modify description: "Create the Score Display UI to track player progress (Progression System - Development & Accomplishment)."
            5. If no relevant gamification element fits naturally, make minimal or no changes.
            6. Output ONLY the updated task JSON object, maintaining the original structure (number, description, explanation, steps, expected_object, recommended_sprites). Do not add any introductory text or explanations outside the JSON. Ensure the output is valid JSON.

            Example Output Format (JSON only):
            {{
                "number": {original_task['number']},
                "description": "Enhanced description incorporating gamification [Expected time...]",
                "explanation": "Enhanced explanation linking to core drives or player types.",
                "recommended_sprites": {json.dumps(original_task.get('recommended_sprites', []))},
                "expected_object": {json.dumps(original_task.get('expected_object', {}))},
                "steps": [
                    "Original step 1",
                    "New step incorporating gamification element (tie to Core Drive)",
                    "Original step 2"
                ]
            }}

            Now, enhance the following task:
            {json.dumps(original_task, indent=2)}
            """

            try:
                print(f"{CYAN}  - Enhancing Chapter {chap_idx+1}, Task {task_idx+1}...{RESET}")
                response_text = await self.generate_initial(prompt)

                json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL | re.IGNORECASE)
                if not json_match:
                    json_match = re.search(r'(\{.*?\})', response_text, re.DOTALL)

                if json_match:
                    json_str = json_match.group(1).strip()
                    try:
                        updated_task_json = json.loads(json_str)
                        if isinstance(updated_task_json, dict) and all(k in updated_task_json for k in ['number', 'description', 'steps']):
                             original_keys = set(original_task.keys())
                             for key in original_keys:
                                 if key not in updated_task_json:
                                     updated_task_json[key] = original_task[key]
                             self.learning_path['chapters'][chap_idx]['tasks'][task_idx] = updated_task_json
                             updated_tasks_count += 1
                             print(f"{GREEN}    ✓ Task enhanced.{RESET}")
                        else:
                             print(f"{YELLOW}    ✗ Enhancement response skipped (missing essential keys or invalid structure).{RESET}")
                             self.metrics_logger.log_error(f"Gamification enhance skip (bad structure): C{chap_idx+1} T{task_idx+1}")
                    except json.JSONDecodeError as json_err:
                        print(f"{RED}    ✗ Error parsing enhancement response for C{chap_idx+1} T{task_idx+1}: {json_err}{RESET}")
                        print(f"{BLUE}      Attempted JSON String: {json_str[:200]}...{RESET}")
                        self.metrics_logger.log_error(f"Gamification enhance JSON parse error: C{chap_idx+1} T{task_idx+1} - {json_err}")
                else:
                    print(f"{YELLOW}    ✗ No JSON found in enhancement response for C{chap_idx+1} T{task_idx+1}. Skipping.{RESET}")
                    print(f"{BLUE}      Raw Response: {response_text[:200]}...{RESET}")
                    self.metrics_logger.log_error(f"Gamification enhance skip (no JSON): C{chap_idx+1} T{task_idx+1}")
            except Exception as e:
                print(f"{RED}    ✗ Unexpected error enhancing C{chap_idx+1} T{task_idx+1}: {e}{RESET}")
                self.metrics_logger.log_error(f"Gamification enhance failed: C{chap_idx+1} T{task_idx+1} - {e}")

            await asyncio.sleep(api_call_delay)

        # Update the metrics logger with the potentially modified learning path
        self.metrics_logger.metrics['learning_path'] = self.learning_path

        print(f"\n{GREEN}Gamification enhancement complete. {updated_tasks_count} tasks updated.{RESET}")
        self.metrics_logger.log_event("gamification_enhancement_finish", details={"updated_count": updated_tasks_count})
        self.save_progress() # Save progress after enhancement

    # --- Coin, Script, and Progress Functions ---
    def display_coin_status(self):
        print(f"\n{CYAN}=========== COIN BALANCE ==========={RESET}")
        print(f"{GREEN}You have {self.user_coins} coins{RESET}")
        print(f"{BLUE}Available actions:{RESET}")
        print(f"- Complete tasks to earn {self.coin_rewards['task_completion']} coins")
        print(f"- Complete chapters to earn {self.coin_rewards['chapter_completion']} additional coins")
        print(f"- Purchase full script implementation for {self.coin_costs['full_script']} coins")
        print(f"- Get a Bonus for 60 coins (does nothing)")
        print(f"- Get a Penalty for 50 coins (does nothing)")
        print(f"{CYAN}======================================{RESET}\n")

    def award_coins(self, reason: str):
        try:
            if reason in self.coin_rewards:
                coins_earned = self.coin_rewards[reason]
                self.user_coins += coins_earned
                self.metrics_logger.log_coin_transaction(coins_earned, reason, self.user_coins)
                print(f"{GREEN}+{coins_earned} coins earned for {reason.replace('_', ' ')}!{RESET}")
                self.save_progress()
                self.display_coin_status()
        except Exception as e:
            self.metrics_logger.log_error(f"Coin award error: {str(e)}")
            print(f"{RED}Error awarding coins: {e}{RESET}")

    def purchase_full_script(self) -> bool:
        try:
            cost = self.coin_costs["full_script"]
            if self.user_coins >= cost:
                self.user_coins -= cost
                self.script_detail_level = "Give the whole script"
                self.metrics_logger.log_coin_transaction(-cost, "script_purchase", self.user_coins)
                self.metrics_logger.log_script_purchase()
                print(f"{GREEN}Purchase successful! {cost} coins deducted.{RESET}")
                print(f"{BLUE}You will now receive the complete script implementation.{RESET}")
                self.save_progress()
                return True
            else:
                print(f"{RED}Not enough coins! You need {cost} coins but have {self.user_coins}.{RESET}")
                return False
        except Exception as e:
            self.metrics_logger.log_error(f"Script purchase error: {str(e)}")
            print(f"{RED}Error during script purchase: {e}{RESET}")
            return False

    def check_script_task(self, task: Dict) -> bool:
        script_keywords = [
            'script', 'coding', 'programming', 'C#', 'csharp',
            'function', 'method', 'class', 'component', 'behavior'
        ]
        task_text = (task.get('description', '') + ' ' +
                    task.get('explanation', '') + ' ' +
                    ' '.join(task.get('steps', []))).lower()
        if any(keyword in task_text for keyword in script_keywords):
            return True

        expected_object = task.get('expected_object', {})
        if isinstance(expected_object, dict):
            components = expected_object.get('components', [])
            if isinstance(components, list):
                for component in components:
                    if isinstance(component, str) and any(kw in component.lower() for kw in ['script', 'behavior']):
                        return True
                    elif isinstance(component, dict) and 'type' in component and any(kw in component['type'].lower() for kw in ['script', 'behavior']):
                         return True
        return False

    def get_existing_scripts(self) -> str:
        scripts_content = ""
        script_count = 0
        scripts_path = os.path.join(self.project_path, 'Assets', 'Scripts')
        alt_paths = [os.path.join(self.project_path, 'Assets')]
        found_scripts = False

        search_paths = [scripts_path] + alt_paths if os.path.exists(scripts_path) else alt_paths

        for search_path in search_paths:
            if os.path.exists(search_path):
                for root, _, files in os.walk(search_path):
                    for file in files:
                        if file.endswith('.cs'):
                            script_path = os.path.join(root, file)
                            try:
                                with open(script_path, 'r', encoding='utf-8') as script_file:
                                    content = script_file.read()
                                    relative_path = os.path.relpath(script_path, self.project_path)
                                    scripts_content += f"\n--- {relative_path} ---\n{content}\n\n"
                                    script_count += 1
                                    found_scripts = True
                            except Exception as e:
                                scripts_content += f"\nError reading {file}: {e}\n"
                if found_scripts and search_path == scripts_path: # Prioritize Scripts folder
                    break

        if script_count == 0:
            return "No C# scripts found in the project."
        return scripts_content

    @staticmethod
    def extract_json_from_response(response: str) -> Dict | None:
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response, re.DOTALL | re.IGNORECASE)
        if json_match:
            json_str = json_match.group(1).strip()
            try:
                return json.loads(json_str)
            except json.JSONDecodeError as e:
                print(f"{RED}JSON Error in markdown block: {e}{RESET}")

        start_index = response.find('{')
        end_index = response.rfind('}')
        if start_index != -1 and end_index != -1 and end_index > start_index:
            json_str = response[start_index:end_index + 1]
            try:
                return json.loads(json_str)
            except json.JSONDecodeError as e:
                print(f"{RED}JSON Error (fallback): {e}{RESET}")
                print(f"{BLUE}Raw Response Snippet:{RESET}\n{response[:500]}...")
                return None
        else:
            print(f"{RED}Could not find valid JSON structure in response.{RESET}")
            print(f"{BLUE}Raw Response Snippet:{RESET}\n{response[:500]}...")
            return None

    def save_progress(self):
        progress_data = {
            'learning_path': self.learning_path,
            'current_chapter_index': self.current_chapter_index,
            'current_task_index': self.current_task_index,
            'user_coins': self.user_coins,
            'script_detail_level': self.script_detail_level
        }
        try:
            # Update metrics with the current learning path *before* saving metrics
            self.metrics_logger.metrics['learning_path'] = self.learning_path
            # Save the progress file
            with open(self.progress_file, 'w') as f:
                json.dump(progress_data, f, indent=4)
            # Now save the metrics file
            self.metrics_logger._save_metrics()
        except Exception as e:
            print(f"{RED}Error saving progress or metrics: {e}{RESET}")
            self.metrics_logger.metrics['errors'].append({
                'timestamp': datetime.datetime.now().isoformat(),
                'error': f"Progress/Metrics Save Error: {e}"
            })
            try:
                 self.metrics_logger._save_metrics()
            except Exception as final_save_e:
                 print(f"{RED}Final attempt to save metrics also failed: {final_save_e}{RESET}")


    def choose_start_mode(self):
        while True:
            print("\n--- Unity Learning Tutor ---")
            print("1. Start a New Learning Path")
            print("2. Continue Previous Progress")
            print("3. Exit")
            choice = input("Enter your choice (1-3): ").strip()

            if choice == '1':
                return 'new'
            elif choice == '2':
                if self.learning_path and 'chapters' in self.learning_path:
                    chapter = self.current_chapter_index + 1
                    task = self.current_task_index + 1
                    print(f"\n{CYAN}Previous Progress Found:{RESET}")
                    print(f"Currently at Chapter {chapter}, Task {task}")
                    return 'continue'
                else:
                    print(f"{YELLOW}No valid previous progress found or learning path missing.{RESET}")
                    print(f"{BLUE}Starting a new learning path.{RESET}")
                    self.learning_path = {}
                    self.current_chapter_index = 0
                    self.current_task_index = 0
                    self.user_coins = 0
                    self.script_detail_level = "Give only the functions and description of each one"
                    self.metrics_logger.metrics['learning_path'] = {}
                    return 'new'
            elif choice == '3':
                print(f"{YELLOW}Exiting tutor.{RESET}")
                sys.exit(0)
            else:
                print(f"{RED}Invalid choice. Please enter 1, 2, or 3.{RESET}")

    def load_progress(self):
        if os.path.exists(self.progress_file):
            try:
                with open(self.progress_file, 'r') as f:
                    progress_data = json.load(f)

                required_keys = ['learning_path', 'current_chapter_index', 'current_task_index']
                if not all(key in progress_data for key in required_keys):
                    print(f"{RED}Incomplete progress data found in {self.progress_file}. Cannot load.{RESET}")
                    self.metrics_logger.metrics['learning_path'] = {}
                    return None

                if not isinstance(progress_data.get('learning_path'), dict) or \
                   not isinstance(progress_data.get('current_chapter_index'), int) or \
                   not isinstance(progress_data.get('current_task_index'), int) or \
                   progress_data.get('current_chapter_index', -1) < 0 or \
                   progress_data.get('current_task_index', -1) < 0:
                    print(f"{RED}Invalid progress data format in {self.progress_file}. Cannot load.{RESET}")
                    self.metrics_logger.metrics['learning_path'] = {}
                    return None

                self.user_coins = progress_data.get('user_coins', 0)
                self.script_detail_level = progress_data.get('script_detail_level', "Give only the functions and description of each one")

                self.learning_path = progress_data['learning_path']
                self.current_chapter_index = progress_data['current_chapter_index']
                self.current_task_index = progress_data['current_task_index']

                # Update the metrics logger with the loaded learning path
                self.metrics_logger.metrics['learning_path'] = self.learning_path

                print(f"{GREEN}Previous progress loaded successfully.{RESET}")
                # Save metrics now that learning path is loaded into it
                self.metrics_logger._save_metrics()
                return progress_data

            except json.JSONDecodeError as e:
                print(f"{RED}Error decoding progress file ({self.progress_file}): {e}{RESET}")
                self.metrics_logger.log_error(f"Progress Load JSON Error: {e}")
                self.metrics_logger.metrics['learning_path'] = {}
                return None
            except Exception as e:
                print(f"{RED}Error loading progress: {e}{RESET}")
                self.metrics_logger.log_error(f"Progress Load Error: {e}")
                self.metrics_logger.metrics['learning_path'] = {}
                return None
        else:
            self.metrics_logger.metrics['learning_path'] = {}
            return None

    # --- Directory Structure and Sprite Handling ---
    def get_directory_structure(self, project_path: str = None) -> Dict[str, Any]:
        if project_path is None:
            project_path = self.project_path
        ui_resource_paths = [
            os.path.join(project_path, 'Assets', 'Sprites'),
            os.path.join(project_path, 'Assets', 'Textures'),
            os.path.join(project_path, 'Assets', 'UI'),
            os.path.join(project_path, 'Assets', 'Resources'),
            os.path.join(project_path, 'Assets')
        ]
        ui_structure = {'directory_text': "", 'sprites': []}
        processed_dirs = set()

        def traverse_directory(directory: str, indent: str = "") -> None:
            norm_dir = os.path.normpath(directory)
            if norm_dir in processed_dirs or not os.path.exists(norm_dir) or not os.path.isdir(norm_dir):
                return
            processed_dirs.add(norm_dir)
            try:
                relative_dir = os.path.relpath(norm_dir, project_path)
                if relative_dir.startswith('Assets'):
                     ui_structure['directory_text'] += f"{indent}|----->{os.path.basename(norm_dir)}\\\n"
                for item in os.listdir(norm_dir):
                    item_path = os.path.join(norm_dir, item)
                    try:
                        if os.path.isdir(item_path):
                            traverse_directory(item_path, indent + "      ")
                        elif os.path.isfile(item_path) and item.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tga', '.psd', '.tiff')):
                            relative_item_path = os.path.relpath(item_path, project_path)
                            if relative_item_path.startswith('Assets'):
                                ui_structure['directory_text'] += f"{indent}      |------>{item}\n"
                                sprite_info = {'filename': item, 'full_path': item_path, 'relative_path': relative_item_path}
                                if not any(s['relative_path'] == relative_item_path for s in ui_structure['sprites']):
                                    ui_structure['sprites'].append(sprite_info)
                    except OSError as oe:
                         print(f"{YELLOW}Warning: Could not access {item_path}: {oe}{RESET}")
                    except Exception as e_inner:
                         print(f"{RED}Error processing item {item_path}: {e_inner}{RESET}")
            except OSError as oe:
                 print(f"{YELLOW}Warning: Could not list directory {norm_dir}: {oe}{RESET}")
            except Exception as e:
                print(f"{RED}Error traversing directory {norm_dir}: {e}{RESET}")

        for potential_path in ui_resource_paths:
            traverse_directory(potential_path)
        ui_structure['sprites'].sort(key=lambda x: x['relative_path'])
        return ui_structure

    def prepare_sprite_details_for_gemini(self, sprite_structure: Dict[str, Any]) -> str:
        if not sprite_structure or not sprite_structure.get('sprites'):
            return "No UI sprites found in the project. Please ensure sprites are placed within the 'Assets' folder (e.g., 'Assets/Sprites', 'Assets/UI')."

        sprite_details = "Available UI Sprites (limit 50 shown):\n------------------------------------\n"
        sprite_limit = 50
        for i, sprite in enumerate(sprite_structure['sprites']):
            if i >= sprite_limit:
                sprite_details += f"... and {len(sprite_structure['sprites']) - sprite_limit} more.\n"
                break
            sprite_details += f"- Filename: {sprite['filename']}\n  Relative Path: {sprite['relative_path'].replace(chr(92), '/')}\n"
        return sprite_details

    # --- Learning Path Generation ---
    async def generate_learning_path(self) -> Dict | None:
        try:
            print(f"{BLUE}Analyzing project structure for UI sprites...{RESET}")
            sprite_structure = self.get_directory_structure()
            sprite_details = self.prepare_sprite_details_for_gemini(sprite_structure)
            print(f"{CYAN}Sprite analysis complete. Found {len(sprite_structure.get('sprites',[]))} sprites.{RESET}")

            game_concept = f"""
            Game Concept for Dyslexia Support Educational Game
            Game Design Document Insights:
            --- GDD START ---
            {self.gdd_content[:3000]}
            --- GDD END ---
            Target Audience: Children with dyslexia (approx. ages 6-10)
            Core Learning Goals: Address challenges in vowel/consonant recognition, phonemic awareness, sound blending, sight words, and basic word formation.
            Platform: Unity 3D (focus on 2D gameplay mechanics)
            Game Worlds/Chapters (Conceptual):
            1. Vowel Voyage / Vowel Island: Focus on vowels.
            2. Consonant Corner / Consonant Caves: Introduce consonants.
            3. Blending Bridge / Blend Bay: Practice blending.
            4. Word Woods / Sentence Shore: Form words/sentences.
            Key Features: Dyslexia-friendly fonts, audio feedback, visual cues, positive reinforcement, gamification.
            Objective: Create an interactive, educational Unity game for dyslexia support.
            """

            prompt = f"""
            You are an expert Unity Game Development Tutor specializing in educational games.
            Create a comprehensive, step-by-step JSON learning path for building a Unity game for children with dyslexia, based on the provided Game Concept and available assets.

            Game Concept Details:
            {game_concept}

            Available UI Sprites in Project:
            {sprite_details}

            Learning Path Requirements:
            1.  Structure: JSON object with "chapters" key (list of 4 chapter objects).
            2.  Chapters: "number" (int), "title" (string), "tasks" (list of task objects).
            3.  Tasks: "number" (int), "description" (string, e.g., "Create Main Menu Scene [Est: 5-10 mins]"), "explanation" (string), "steps" (list of strings), "expected_object" (JSON: 'type', 'name', 'components'/'properties'), "recommended_sprites" (list of strings or descriptive names).
            4.  Content Focus: Project setup, scene creation, UI (main menu), core gameplay mechanics per chapter, player input, scoring, feedback placeholders, scene transitions, basic gamification.
            5.  Practicality: Granular tasks for beginner/intermediate Unity user.
            6.  Scripting: Include C# tasks. Specify filename, high-level overview in steps (NOT full code), update "expected_object".
            7.  Sprite Usage: Reference available sprites by filename accurately. List descriptive ideal filenames if missing.
            8.  Dyslexia Considerations: Mention relevant considerations in tasks (e.g., "Use clear font").
            9.  Output Format: ONLY the valid JSON object. No extra text or markdown.

            Example Task Structure (within chapters list):
            {{
                "number": 1,
                "title": "Chapter 1: Project Setup and Main Menu",
                "tasks": [
                    {{ "number": 1, "description": "Create new Unity 2D Project [Est: 5 mins]", ... }},
                    {{ "number": 2, "description": "Create Main Menu Scene [Est: 5 mins]", ... }}
                ]
            }}
            """

            print(f"{CYAN}Generating learning path with Gemini... This may take a minute or two.{RESET}")
            response_text = await self.generate_initial(prompt)
            learning_path_json = UnityLearningTutor.extract_json_from_response(response_text)

            if learning_path_json and 'chapters' in learning_path_json and isinstance(learning_path_json['chapters'], list):
                print(f"{GREEN}✓ Learning path generated successfully.{RESET}")
                self.learning_path = learning_path_json
                self.metrics_logger.metrics['learning_path'] = self.learning_path
                return self.learning_path
            else:
                print(f"{RED}Error: Failed to generate a valid learning path JSON from the response.{RESET}")
                self.metrics_logger.log_error("Learning path generation failed - invalid JSON structure")
                self.metrics_logger.metrics['learning_path'] = {}
                return None

        except Exception as e:
            print(f"{RED}An unexpected error occurred during learning path generation: {e}{RESET}")
            self.metrics_logger.log_error(f"Learning path generation error: {str(e)}")
            traceback.print_exc()
            self.metrics_logger.metrics['learning_path'] = {}
            return None

    async def generate_initial(self, prompt: str) -> str:
        try:
            response = await asyncio.to_thread(
                self.gemini_client.generate_content,
                prompt,
            )
            if response.parts:
                return response.text
            else:
                block_reason = response.prompt_feedback.block_reason if response.prompt_feedback else "Unknown"
                error_msg = f"Gemini response was empty or blocked. Reason: {block_reason}"
                print(f"{RED}{error_msg}{RESET}")
                self.metrics_logger.log_error(error_msg)
                raise ValueError(error_msg)
        except Exception as e:
            print(f"{RED}Error communicating with Gemini API: {e}{RESET}")
            self.metrics_logger.log_error(f"Gemini API Error: {str(e)}")
            raise

    # --- Task Presentation, Validation, and Watchdog ---
    async def show_coin_options_menu(self, current_task: Dict) -> bool:
        self.metrics_logger.log_menu_access()
        while True:
            print(f"\n{CYAN}=========== COIN MENU ==========={RESET}")
            print(f"{GREEN}You have {self.user_coins} coins{RESET}")
            print(f"{BLUE}Available options:{RESET}")
            print(f"1. Purchase full script implementation ({self.coin_costs['full_script']} coins)")
            print(f"2. Get a Bonus (60 coins - does nothing)")
            print(f"3. Get a Penalty (50 coins - does nothing)")
            print(f"4. Return to task")
            print(f"{CYAN}======================================{RESET}\n")
            choice = input(f"{CYAN}Enter your choice (1-4): {RESET}").strip()

            if choice == '1':
                if self.purchase_full_script():
                    return True
            elif choice == '2':
                cost = 60
                if self.user_coins >= cost:
                    self.user_coins -= cost
                    print(f"{YELLOW}You spent {cost} coins for a Bonus! (This has no game effect){RESET}")
                    self.metrics_logger.log_coin_transaction(-cost, "bonus_purchase", self.user_coins)
                    self.save_progress()
                    self.display_coin_status()
                else:
                    print(f"{RED}Not enough coins! You need {cost} coins but have {self.user_coins}.{RESET}")
            elif choice == '3':
                cost = 50
                if self.user_coins >= cost:
                    self.user_coins -= cost
                    print(f"{YELLOW}You spent {cost} coins for a Penalty! (This has no game effect){RESET}")
                    self.metrics_logger.log_coin_transaction(-cost, "penalty_purchase", self.user_coins)
                    self.save_progress()
                    self.display_coin_status()
                else:
                    print(f"{RED}Not enough coins! You need {cost} coins but have {self.user_coins}.{RESET}")
            elif choice == '4':
                print(f"{BLUE}Returning to task...{RESET}")
                return False
            else:
                print(f"{RED}Invalid choice. Please enter a number between 1 and 4.{RESET}")

    async def validate_task_completion(self, current_task: Dict) -> bool:
        try:
            current_chapter = self.learning_path['chapters'][self.current_chapter_index]
            self.metrics_logger.log_validation_attempt(current_chapter['number'], current_task['number'])

            is_script_task = self.check_script_task(current_task)
            script_content_generated = False

            if is_script_task:
                print(f"\n{MAGENTA}=== SCRIPTING TASK DETAILS ==={RESET}")
                print(f"This task requires creating or modifying a C# script.")
                print(f"Your current script detail level: {YELLOW}{self.script_detail_level}{RESET}")
                print(f"Coins needed for full script: {self.coin_costs['full_script']} (You have: {self.user_coins})")

                if self.script_detail_level != "Give the whole script":
                    while True:
                        choice = input(f"{CYAN}Show script details now? (yes/buy/no): {RESET}").strip().lower()
                        if choice in ['yes', 'y']: break
                        elif choice == 'buy':
                            if self.purchase_full_script(): break
                            else: continue
                        elif choice in ['no', 'n']:
                            print(f"{BLUE}Okay, skipping script generation for now. Remember to implement the required logic.{RESET}")
                            script_content_generated = False
                            break
                        else: print(f"{RED}Invalid input. Please enter 'yes', 'buy', or 'no'.{RESET}")
                    if choice != 'no':
                         script_content = await self.generate_script_content(current_task)
                         print(f"\n{BLUE}--- Generated Script Content ({self.script_detail_level}) ---{RESET}\n{script_content}\n{BLUE}----------------------------------------------------{RESET}")
                         print(f"\n{CYAN}Implement this script logic in your Unity project ('{current_task.get('expected_object',{}).get('name','YourScript.cs')}').{RESET}")
                         script_content_generated = True
                else:
                    script_content = await self.generate_script_content(current_task)
                    print(f"\n{BLUE}--- Generated Script Content ({self.script_detail_level}) ---{RESET}\n{script_content}\n{BLUE}----------------------------------------------------{RESET}")
                    print(f"\n{CYAN}Implement this script logic in your Unity project ('{current_task.get('expected_object',{}).get('name','YourScript.cs')}').{RESET}")
                    script_content_generated = True
                input(f"{CYAN}Press Enter when you are ready to proceed with implementing the task in Unity...{RESET}")

            print(f"\n{GREEN}=== TASK VALIDATION PENDING ==={RESET}")
            print(f"Now, please perform the steps described for this task in the Unity Editor.")
            print(f"{CYAN}I will check your work automatically when you SAVE THE SCENE in Unity.{RESET}")
            print(f"{YELLOW}Type 'menu' and press Enter anytime to access coin options or 'skip' to bypass validation (not recommended).{RESET}")

            self._scene_change_flag.clear()
            start_time = time.time()
            timeout_seconds = 600
            validation_triggered = False

            while not validation_triggered:
                if time.time() - start_time > timeout_seconds:
                    print(f"\n{RED}Validation timed out after {timeout_seconds} seconds.{RESET}")
                    print(f"{YELLOW}Please ensure you have saved the scene in Unity.{RESET}")
                    retry = input(f"{CYAN}Do you want to wait longer? (yes/no): {RESET}").strip().lower()
                    if retry in ['yes', 'y']: start_time = time.time(); continue
                    else: self.metrics_logger.log_event("validation_timeout"); return False

                if self._scene_change_flag.is_set():
                    print(f"\n{GREEN}Scene save detected! Analyzing changes...{RESET}")
                    self._scene_change_flag.clear(); validation_triggered = True; break

                user_command = None
                if os.name == 'nt':
                    if msvcrt.kbhit():
                        char = msvcrt.getch()
                        try:
                           decoded_char = char.decode('utf-8').lower()
                           if decoded_char in ['m', 's']: user_command = input(f"\nEnter command ('menu' or 'skip'): ").strip().lower()
                        except UnicodeDecodeError: pass
                else:
                    import select
                    rlist, _, _ = select.select([sys.stdin], [], [], 0.1)
                    if rlist: user_command = sys.stdin.readline().strip().lower()

                if user_command:
                    if user_command == 'menu':
                        print(f"{CYAN}Opening coin options menu...{RESET}")
                        self.metrics_logger.log_event("menu_access_during_validation")
                        purchased_full_script = await self.show_coin_options_menu(current_task)
                        if purchased_full_script and is_script_task:
                            print(f"{GREEN}Regenerating script with full implementation...{RESET}")
                            script_content = await self.generate_script_content(current_task)
                            print(f"\n{BLUE}--- FULL SCRIPT IMPLEMENTATION ---{RESET}\n{script_content}\n{BLUE}----------------------------------{RESET}")
                            input(f"{CYAN}Press Enter when ready to continue validation...{RESET}")
                        start_time = time.time()
                        print(f"{CYAN}Resuming validation check. Remember to save the scene in Unity.{RESET}")
                    elif user_command == 'skip':
                        confirm_skip = input(f"{YELLOW}Are you sure you want to skip validation for this task? (yes/no): {RESET}").strip().lower()
                        if confirm_skip in ['yes', 'y']:
                            print(f"{YELLOW}Skipping task validation. Moving to the next task.{RESET}")
                            self.metrics_logger.log_task_completion(current_chapter['number'], current_task['number'], False)
                            self.metrics_logger.log_event("validation_skipped_by_user", {"chapter": current_chapter['number'], "task": current_task['number']})
                            return True
                        else: print(f"{BLUE}Resuming validation check.{RESET}"); start_time = time.time()
                    else: print(f"{RED}Unknown command '{user_command}'. Type 'menu' or 'skip'.{RESET}")

                await asyncio.sleep(0.2)

            print(f"{BLUE}Performing scene analysis...{RESET}")
            scene_analysis = comprehensive_scene_analysis(self.project_path)
            self.metrics_logger.log_scene_analysis()

            validation_prompt = f"""
            You are a Unity Task Validator AI. Analyze the provided Unity scene data and determine if the user has successfully completed the specified task.

            Task Details:
            - Chapter: {current_chapter.get('number', 'N/A')} - {current_chapter.get('title', 'N/A')}
            - Task: {current_task.get('number', 'N/A')} - {current_task.get('description', 'N/A')}
            - Explanation: {current_task.get('explanation', 'N/A')}
            - Steps Overview: {"; ".join(current_task.get('steps', []))}
            - Expected Outcome (Key GameObject/Component/Asset):
            ```json
            {json.dumps(current_task.get('expected_object', {}), indent=2)}
            ```

            Current Unity Scene Analysis Data:
            ```json
            {json.dumps(scene_analysis, indent=2)}
            ```

            Instructions:
            1.  Compare 'Current Unity Scene Analysis Data' against 'Expected Outcome'.
            2.  Check for primary GameObject/Asset by name and type.
            3.  Verify required components/properties.
            4.  Consider task steps - does the scene reflect actions?
            5.  Provide clear feedback, especially if incomplete. List missing elements.
            6.  Respond ONLY with valid JSON: {{ "completed": boolean, "feedback": string, "missing_components": list[string] }}.

            Example JSON Response (Incomplete):
            {{ "completed": false, "feedback": "Good start! 'Player' GameObject exists, but 'PlayerMovement' script is missing.", "missing_components": ["PlayerMovement script component on Player GameObject"] }}

            Example JSON Response (Complete):
            {{ "completed": true, "feedback": "Excellent work! 'MainMenu' scene contains expected elements.", "missing_components": [] }}

            Validate the task. Output ONLY the JSON object.
            """

            print(f"{BLUE}Sending analysis to Gemini for validation...{RESET}")
            validation_response_text = await self.generate_initial(validation_prompt)
            validation_result = UnityLearningTutor.extract_json_from_response(validation_response_text)

            if validation_result and isinstance(validation_result, dict) and 'completed' in validation_result and 'feedback' in validation_result:
                print(f"\n{MAGENTA}--- VALIDATION RESULT ---{RESET}")
                print(f"{CYAN}Feedback:{RESET} {validation_result.get('feedback', 'No feedback provided.')}")

                if validation_result.get('completed', False):
                    print(f"{GREEN}✓ Task Completed Successfully!{RESET}")
                    self.save_progress()
                    return True
                else:
                    print(f"{RED}✗ Task Not Yet Complete.{RESET}")
                    missing = validation_result.get('missing_components', [])
                    if missing:
                        print(f"{YELLOW}Missing/Incorrect Elements:{RESET}")
                        for item in missing: print(f"  - {item}")
                    print(f"{CYAN}Please review feedback, make changes in Unity, and save the scene again.{RESET}")
                    self.metrics_logger.log_task_completion(current_chapter['number'], current_task['number'], False)
                    return False
            else:
                print(f"{RED}Validation Error: Could not parse valid result from AI.{RESET}")
                print(f"{BLUE}Raw Response Snippet:{RESET}\n{validation_response_text[:500]}...")
                self.metrics_logger.log_error("Validation response parsing error")
                print(f"{YELLOW}Skipping validation attempt. Please try saving again.{RESET}")
                return False

        except Exception as e:
            print(f"{RED}An unexpected error occurred during task validation: {e}{RESET}")
            self.metrics_logger.log_error(f"Task validation failed: {str(e)}")
            traceback.print_exc()
            return False

    async def wait_for_scene_change(self):
        while not self._scene_change_flag.is_set():
            await asyncio.sleep(0.5)
        self._scene_change_flag.clear()

    async def present_current_task(self) -> bool:
        if not self.learning_path or not self.learning_path.get('chapters'):
            print(f"{RED}Error: Learning path is missing or invalid.{RESET}")
            return False

        if self.current_chapter_index >= len(self.learning_path['chapters']):
            os.system('cls' if os.name == 'nt' else 'clear')
            print(f"{GREEN}=========================\n🎉 TUTORIAL COMPLETE! 🎉\n========================={RESET}")
            print(f"\n{CYAN}Congratulations! You've finished all chapters.{RESET}")
            print(f"You have {self.user_coins} coins remaining.")
            print(f"Check 'learning_metrics.md' for a session summary.")
            self.metrics_logger.log_event("tutorial_completed")
            return False

        current_chapter = self.learning_path['chapters'][self.current_chapter_index]

        if self.current_task_index >= len(current_chapter.get('tasks', [])):
            print(f"\n{GREEN}Chapter {current_chapter.get('number', self.current_chapter_index + 1)} completed!{RESET}")
            if self.current_chapter_index < len(self.learning_path['chapters']):
                 self.award_coins('chapter_completion')
                 self.metrics_logger.log_chapter_completion(current_chapter['number'])

            self.current_chapter_index += 1
            self.current_task_index = 0
            self.save_progress()

            if self.current_chapter_index >= len(self.learning_path['chapters']):
                 return await self.present_current_task()

            next_chapter = self.learning_path['chapters'][self.current_chapter_index]
            self.metrics_logger.log_chapter_start(next_chapter['number'], next_chapter.get('title', 'Untitled Chapter'))
            return await self.present_current_task()

        current_task = current_chapter['tasks'][self.current_task_index]
        os.system('cls' if os.name == 'nt' else 'clear')
        print(f"{CYAN}=== UNITY LEARNING TUTOR ==={RESET}")
        self.display_coin_status()

        print(f"\n{BLUE}CHAPTER {current_chapter.get('number', self.current_chapter_index + 1)}: {current_chapter.get('title', 'Untitled Chapter')}{RESET}")
        print(f"{GREEN}TASK {current_task.get('number', self.current_task_index + 1)}: {current_task.get('description', 'No Description')}{RESET}")
        print(f"\n{YELLOW}Explanation:{RESET}\n{textwrap.fill(current_task.get('explanation', 'N/A'), width=80)}")
        print(f"\n{CYAN}Steps to Complete:{RESET}")
        steps = current_task.get('steps', [])
        if steps:
            for i, step in enumerate(steps, 1): print(f"{i}. {textwrap.fill(step, width=78, initial_indent='   ', subsequent_indent='   ')}")
        else: print("   No specific steps provided.")

        if self.check_script_task(current_task):
            print(f"\n{MAGENTA}NOTE: This task involves C# scripting! Details during validation.{RESET}")

        recommended_sprites = current_task.get('recommended_sprites', [])
        if recommended_sprites:
            print(f"\n{YELLOW}Recommended Sprites/Assets:{RESET}")
            for sprite_name in recommended_sprites: print(f"  - {sprite_name}")
            print(f"{BLUE}(Check Assets or create/import similar assets){RESET}")

        print(f"\n{GREEN}---> Implement these changes in the Unity Editor.{RESET}")
        self.metrics_logger.log_task_start(current_chapter['number'], current_task['number'], current_task.get('description', 'Untitled Task'))
        return True

    # --- MODIFIED run_learning_path ---
    async def run_learning_path(self):
        """Main loop to present tasks and handle validation."""
        try:
            start_mode = self.choose_start_mode()
            apply_enhancements = False

            if start_mode == 'new':
                print(f"{CYAN}Starting a new learning path...{RESET}")
                self.learning_path = {}
                self.current_chapter_index = 0
                self.current_task_index = 0
                self.user_coins = 0
                self.script_detail_level = "Give only the functions and description of each one"
                self.metrics_logger.metrics['learning_path'] = {}
                generated_path = await self.generate_learning_path()
                if not generated_path:
                    print(f"{RED}Failed to generate learning path. Exiting.{RESET}")
                    return
                self.save_progress()
                self.metrics_logger.log_event("new_learning_path_generated")

            elif start_mode == 'continue':
                if not self.learning_path or not self.learning_path.get('chapters'):
                    print(f"{RED}Error: Continue mode selected, but no valid learning path found.{RESET}")
                    return
                print(f"{GREEN}Resuming previous learning path...{RESET}")
                self.metrics_logger.log_event("resumed_learning_path")

                # --- MODIFICATION START: Always ask about enhancement when continuing ---
                if self.gdd_was_validated:
                    print(f"{YELLOW}NOTICE: Your Game Design Document file seems to have changed since you last saved progress.{RESET}")
                else:
                    print(f"{BLUE}Continuing progress. You can choose to enhance the path based on the last GDD analysis.{RESET}")

                # --- ASK THE USER (Always) ---
                while True:
                    choice = input(f"{CYAN}Would you like to enhance the remaining learning path based on the GDD analysis? By selecting yes the undone tasks will change , by selecting no nothing changed and the previouse tasks will remaine the same (yes/no): {RESET}").strip().lower()
                    if choice in ['yes', 'y']:
                        if self.validation_results: # Check if we actually have results
                            print(f"{BLUE}Applying enhancements based on the GDD analysis...{RESET}")
                            apply_enhancements = True
                            self.metrics_logger.log_event("gdd_enhancement_accepted_on_continue")
                        else:
                            print(f"{YELLOW}Cannot enhance as GDD validation results are missing. Continuing without enhancements.{RESET}")
                            self.metrics_logger.log_event("gdd_enhancement_failed_no_results_on_continue")
                        break # Exit the asking loop
                    elif choice in ['no', 'n']:
                        print(f"{BLUE}Okay, continuing with the existing learning path structure without enhancements.{RESET}")
                        apply_enhancements = False
                        self.metrics_logger.log_event("gdd_enhancement_declined_on_continue")
                        break # Exit the asking loop
                    else:
                        print(f"{RED}Invalid input. Please enter 'yes' or 'no'.{RESET}")
                # --- END ASK THE USER ---
                # --- MODIFICATION END ---

                # Apply enhancements if flagged
                if apply_enhancements and self.validation_results:
                     await self.enhance_learning_path_with_gamification(self.validation_results)

            self.setup_watchdog()
            self.metrics_logger._save_metrics() # Ensure metrics are saved with current path

            if self.current_task_index == 0 and self.current_chapter_index < len(self.learning_path.get('chapters', [])):
                 first_chapter = self.learning_path['chapters'][self.current_chapter_index]
                 first_chapter_key = f"chapter_{first_chapter['number']}"
                 if first_chapter_key not in self.metrics_logger.metrics['chapters']:
                     self.metrics_logger.log_chapter_start(first_chapter['number'], first_chapter.get('title', 'Untitled Chapter'))

            while await self.present_current_task():
                if self.current_chapter_index >= len(self.learning_path.get('chapters', [])): break
                current_chapter = self.learning_path['chapters'][self.current_chapter_index]
                if self.current_task_index >= len(current_chapter.get('tasks',[])):
                     print(f"{YELLOW}Warning: Reached end of tasks unexpectedly. Advancing chapter.{RESET}")
                     continue

                current_task = current_chapter['tasks'][self.current_task_index]
                task_completed = await self.validate_task_completion(current_task)

                if task_completed:
                    self.metrics_logger.log_task_completion(current_chapter['number'], current_task['number'], True)
                    print(f"{GREEN}Task validated successfully!{RESET}")
                    self.award_coins('task_completion')
                    input(f"\n{CYAN}Press Enter to continue to the next task...{RESET}")
                    self.current_task_index += 1
                    self.save_progress()
                else:
                    print(f"{YELLOW}Task validation was not successful.{RESET}")
                    print(f"{BLUE}Review steps/feedback, make corrections, and save scene again.{RESET}")
                    input(f"{CYAN}Press Enter to re-display the current task instructions...{RESET}")

        except asyncio.CancelledError:
             print(f"\n{YELLOW}Learning path execution cancelled.{RESET}")
             self.metrics_logger.log_event("learning_path_cancelled")
        except Exception as e:
            print(f"{RED}An error occurred during the learning path execution: {e}{RESET}")
            self.metrics_logger.log_error(f"Learning path run error: {str(e)}")
            traceback.print_exc()
        finally:
            print(f"{CYAN}Ending learning session.{RESET}")
            if hasattr(self, 'metrics_logger') and self.metrics_logger: # Check if logger exists
                self.metrics_logger.log_session_end()
            self.cleanup_watchdog()
            print(f"{BLUE}Session metrics saved to 'learning_metrics.md'.{RESET}")


    async def generate_script_content(self, task: Dict) -> str:
        try:
            print(f"{BLUE}Fetching existing script context...{RESET}")
            existing_scripts = self.get_existing_scripts()
            expected_obj = task.get('expected_object', {})
            script_name = expected_obj.get('name', 'YourScript.cs')
            if not script_name.endswith(".cs"): script_name += ".cs"

            prompt = f"""
            You are an expert Unity C# Scripting Assistant. Generate C# script content for the task, considering detail level and existing scripts.

            Task Context:
            - Description: {task.get('description', 'N/A')}
            - Explanation: {task.get('explanation', 'N/A')}
            - Key Steps: {"; ".join(s for s in task.get('steps', []) if any(kw in s.lower() for kw in ['script', 'code', 'function', 'variable', 'logic']))}
            - Expected Script Name: {script_name}
            - Expected Association: {json.dumps(expected_obj, indent=2)}

            Requested Detail Level: "{self.script_detail_level}"

            Existing Scripts Context (truncated):
            --- START EXISTING SCRIPTS ---
            {existing_scripts[:4000]}
            --- END EXISTING SCRIPTS ---

            Instructions:
            1. Generate complete C# script for '{script_name}'.
            2. Adhere strictly to 'Requested Script Detail Level':
               * "Give only the functions and description": Skeleton script with signatures and comments (// TODO: or /// <summary>). NO implementation logic.
               * "Give the whole script": Full, compilable implementation with logic.
            3. Include standard Unity namespaces.
            4. Define class correctly (public class {script_name.replace('.cs','')} : MonoBehaviour).
            5. Implement logic based on task details. Make reasonable assumptions if vague.
            6. Consider interactions with other scripts.
            7. Format output as C# code within ```csharp ... ``` markers. No explanations outside code block.

            Generate the script content:
            """

            print(f"{CYAN}Generating script content ({self.script_detail_level})...{RESET}")
            script_response = await self.generate_initial(prompt)
            code_block_match = re.search(r'```csharp\s*(.*?)\s*```', script_response, re.DOTALL | re.IGNORECASE)

            if code_block_match:
                extracted_code = code_block_match.group(1).strip()
                if "public class" in extracted_code and "MonoBehaviour" in extracted_code:
                     print(f"{GREEN}✓ Script content generated.{RESET}")
                     return extracted_code
                else:
                     print(f"{YELLOW}Warning: Generated content might not be complete C# script.{RESET}")
                     return extracted_code
            else:
                 print(f"{YELLOW}Warning: Could not find ```csharp block. Displaying raw response:{RESET}")
                 self.metrics_logger.log_error(f"Script generation failed - no C# block: {task.get('description')}")
                 return script_response

        except Exception as e:
            print(f"{RED}Error generating script content: {e}{RESET}")
            self.metrics_logger.log_error(f"Script generation error: {str(e)}")
            traceback.print_exc()
            return f"// Error generating script: {e}"

    def setup_watchdog(self):
        try:
            class ProjectWatchdog(FileSystemEventHandler):
                def __init__(self, learning_tutor_instance):
                    self.learning_tutor = learning_tutor_instance
                    self.last_event_times = {}
                    self.debounce_interval = 1.5

                def on_modified(self, event: FileSystemEvent):
                    if event.is_directory or not event.src_path.lower().endswith('.unity'): return
                    current_time = time.time()
                    file_path = event.src_path
                    last_time = self.last_event_times.get(file_path, 0)
                    if current_time - last_time > self.debounce_interval:
                        self.last_event_times[file_path] = current_time
                        print(f"{MAGENTA}Watchdog: Detected modification in {os.path.basename(file_path)}{RESET}")
                        try:
                            loop = asyncio.get_running_loop()
                            asyncio.run_coroutine_threadsafe(self.set_scene_change_flag(), loop)
                        except RuntimeError:
                            print(f"{YELLOW}Watchdog: No running event loop found.{RESET}")

                async def set_scene_change_flag(self):
                    print(f"{GREEN}Watchdog: Signaling scene change flag...{RESET}")
                    self.learning_tutor._scene_change_flag.set()

            if not os.path.isdir(self.project_path):
                 print(f"{RED}Error: Project path '{self.project_path}' invalid. Cannot start watchdog.{RESET}")
                 self.metrics_logger.log_error(f"Watchdog setup failed: Invalid project path '{self.project_path}'")
                 return

            self.watchdog = ProjectWatchdog(self)
            self.watchdog_observer = Observer()
            watch_path = os.path.join(self.project_path, 'Assets')
            if not os.path.isdir(watch_path):
                 print(f"{YELLOW}Warning: 'Assets' dir not found. Watching project root.{RESET}")
                 watch_path = self.project_path

            self.watchdog_observer.schedule(self.watchdog, watch_path, recursive=True)
            self.watchdog_observer.start()
            print(f"{BLUE}Watchdog started monitoring: {watch_path}{RESET}")
            self.metrics_logger.log_event("watchdog_started", {"path": watch_path})

        except Exception as e:
            print(f"{RED}Error setting up watchdog: {e}{RESET}")
            self.metrics_logger.log_error(f"Watchdog setup error: {str(e)}")
            traceback.print_exc()

    def cleanup_watchdog(self):
        try:
            if self.watchdog_observer and self.watchdog_observer.is_alive():
                self.watchdog_observer.stop()
                self.watchdog_observer.join(timeout=2)
                print(f"{BLUE}Watchdog stopped.{RESET}")
                self.metrics_logger.log_event("watchdog_stopped")
            self.watchdog_observer = None
            self.watchdog = None
        except Exception as e:
            print(f"{RED}Error cleaning up watchdog: {e}{RESET}")
            if hasattr(self, 'metrics_logger') and self.metrics_logger: # Check if logger exists
                self.metrics_logger.log_error(f"Watchdog cleanup error: {str(e)}")

# --- Main Execution (Modified) ---
async def main():
    os.system('cls' if os.name == 'nt' else 'clear')
    print(f"{GREEN}=== Dyslexia Game Development Tutor ===")
    print(f"{YELLOW}      Version 1.5 (Conditional GDD Validation)      {RESET}\n")
    print(f"{BLUE}PROJECT PURPOSE:{RESET}")
    print(textwrap.fill("Guide creation of a Unity game for dyslexia support.", width=80))
    print(f"\n{CYAN}TIPS FOR SUCCESS:{RESET}")
    tips = [
        "• Keep Unity Editor open.", "• SAVE SCENE in Unity to trigger validation.",
        "• Follow tasks sequentially.", "• Use coins wisely for script help.",
        "• Name assets descriptively.", "• Type 'menu' during validation for options."
    ]
    print('\n'.join(tips)); print("\n" + "="*60)
    input(f"\n{GREEN}Press Enter to begin setup...{RESET}")
    os.system('cls' if os.name == 'nt' else 'clear')

    project_path = None
    while not project_path:
        print(f"{CYAN}Select your Unity Project Folder:{RESET}")
        input(f"{GREEN}Press Enter for folder selection dialog...{RESET}")
        project_path = select_project_path()
        if not project_path:
            print(f"{RED}No project folder selected.{RESET}")
            if input(f"{CYAN}Try again? (yes/no): {RESET}").strip().lower() not in ['yes', 'y']: return
        elif not os.path.exists(os.path.join(project_path, 'Assets')):
             print(f"{RED}Error: Invalid Unity project (missing 'Assets' folder).{RESET}")
             project_path = None
             if input(f"{CYAN}Select different folder? (yes/no): {RESET}").strip().lower() not in ['yes', 'y']: return
        else: print(f"{GREEN}✓ Project location set:{RESET} {project_path}")

    gdd_path = None
    while not gdd_path:
        print(f"\n{CYAN}Select your Game Design Document (GDD):{RESET}")
        input(f"{GREEN}Press Enter for file selection dialog...{RESET}")
        gdd_path = select_gdd_file()
        if not gdd_path:
            print(f"{RED}No GDD file selected.{RESET}")
            if input(f"{CYAN}Try again? (yes/no): {RESET}").strip().lower() not in ['yes', 'y']: return
        else: print(f"{GREEN}✓ GDD selected:{RESET} {gdd_path}")

    print(f"\n{CYAN}Extracting content from GDD...{RESET}")
    gdd_content = extract_gdd_content(gdd_path)
    if not gdd_content:
        print(f"{RED}Error: Could not extract GDD content.{RESET}"); input(f"{GREEN}Press Enter to exit...{RESET}"); return
    print(f"{GREEN}✓ GDD content extracted.{RESET}")

    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        print(f"\n{CYAN}Enter your Google AI (Gemini) API Key:{RESET}")
        gemini_api_key = input(f"{CYAN}API Key: {RESET}").strip()
    if not gemini_api_key: print(f"{RED}API key required. Exiting.{RESET}"); return

    validation_results = None
    gdd_was_validated = False
    gemini_client = None
    temp_metrics_logger = None
    try:
        temp_metrics_logger = MetricsLogger(project_path)
    except Exception as log_init_e:
        print(f"{RED}CRITICAL ERROR initializing metrics logger: {log_init_e}{RESET}")

    try:
        print(f"\n{CYAN}Configuring/testing Gemini API...{RESET}")
        genai.configure(api_key=gemini_api_key)
        gemini_client = genai.GenerativeModel('gemini-2.5-flash-preview-04-17')
        test_response = await asyncio.to_thread(gemini_client.generate_content, "Respond with 'OK'")
        if not test_response or 'ok' not in test_response.text.lower():
            feedback = test_response.prompt_feedback if hasattr(test_response, 'prompt_feedback') else 'N/A'
            raise ConnectionError(f"API key validation failed. Feedback: {feedback}")
        print(f"{GREEN}✓ Gemini API connection successful.{RESET}")

        print(f"\n{CYAN}Checking GDD hash...{RESET}")
        gdd_hash_file_path = os.path.join(project_path, GDD_HASH_FILE)
        current_gdd_hash = calculate_gdd_hash(gdd_content)
        saved_gdd_hash = None
        if os.path.exists(gdd_hash_file_path):
            try:
                with open(gdd_hash_file_path, 'r') as f: saved_gdd_hash = f.read().strip()
            except Exception as e:
                print(f"{RED}Error loading saved GDD hash: {e}{RESET}")
                if temp_metrics_logger: temp_metrics_logger.log_error(f"GDD Hash Load Error: {e}")

        if saved_gdd_hash == current_gdd_hash:
            print(f"{GREEN}GDD unchanged. Loading previous validation results.{RESET}")
            results_path = os.path.join(project_path, VALIDATION_RESULTS_FILE)
            if os.path.exists(results_path):
                try:
                    with open(results_path, 'r') as f: validation_results = json.load(f)
                    print(f"{GREEN}✓ Loaded previous validation results.{RESET}")
                except Exception as e:
                    print(f"{RED}Error loading previous validation results: {e}{RESET}")
                    if temp_metrics_logger: temp_metrics_logger.log_error(f"GDD Results Load Error: {e}")
            else: print(f"{YELLOW}Previous validation results not found.{RESET}")
            gdd_was_validated = False
        else:
            print(f"{YELLOW if saved_gdd_hash else BLUE}{'GDD hash mismatch! Re-analyzing...' if saved_gdd_hash else 'No previous GDD hash. Analyzing...'}{RESET}")
            print(f"{CYAN}Analyzing GDD with Octalysis... (This may take a minute){RESET}")
            validation_results = validate_gdd_with_octalysis(gemini_client, gdd_content)
            gdd_was_validated = True
            if "error" in validation_results: raise ValueError(f"GDD Validation Error: {validation_results['error']}")

            report_path = save_validation_report(validation_results, project_path)
            try:
                with open(gdd_hash_file_path, 'w') as f: f.write(current_gdd_hash)
                print(f"{BLUE}Saved new GDD hash.{RESET}")
            except Exception as e:
                print(f"{RED}Error saving new GDD hash: {e}{RESET}")
                if temp_metrics_logger: temp_metrics_logger.log_error(f"GDD Hash Save Error: {e}")

            print(f"\n{GREEN}=== GDD ANALYSIS COMPLETE ===")
            print(f"{BLUE}Octalysis Score: {validation_results.get('overall_score', 'N/A')}/10{RESET}")
            print(f"{BLUE}Summary: {validation_results.get('summary', 'N/A')}{RESET}")
            print(f"\n{GREEN}Top Recommendations:{RESET}")
            for i, rec in enumerate(validation_results.get('top_recommendations', [])[:3], 1): print(f" {i}. {rec}")
            if report_path: print(f"\n{CYAN}Detailed report saved to: {report_path}{RESET}")
            else: print(f"\n{YELLOW}Could not save detailed report.{RESET}")

    except (ConnectionError, ValueError, Exception) as e:
        error_message = str(e)
        print(f"\n{RED}--- ERROR DURING SETUP ---{RESET}\n{RED}Error: {error_message}{RESET}")
        traceback.print_exc()
        if "api key not valid" in error_message.lower(): print(f"{YELLOW}Suggestion: Check Gemini API key.{RESET}")
        elif "quota" in error_message.lower(): print(f"{YELLOW}Suggestion: Check API quota.{RESET}")
        elif isinstance(e, json.JSONDecodeError): print(f"{YELLOW}Suggestion: AI response JSON error. Try again.{RESET}")
        elif "resource has been exhausted" in error_message.lower(): print(f"{YELLOW}Suggestion: API service busy. Try again later.{RESET}")
        else: print(f"{YELLOW}Unexpected error. Check details.{RESET}")
        print(f"\n{YELLOW}Cannot proceed.{RESET}")
        if temp_metrics_logger:
            temp_metrics_logger.log_error(f"Setup failed: {error_message}")
            temp_metrics_logger.log_session_end()
        input(f"{GREEN}Press Enter to exit...{RESET}")
        return

    tutor = None
    try:
        print(f"\n{GREEN}Initializing Unity Learning Tutor...{RESET}")
        tutor = UnityLearningTutor(gemini_api_key, project_path, gdd_content,
                                   validation_results, gdd_was_validated)
        await tutor.run_learning_path()
    except Exception as e:
        print(f"\n{RED}--- ERROR DURING LEARNING PATH ---{RESET}\n{RED}Error: {e}{RESET}")
        traceback.print_exc()
        print(f"\n{YELLOW}TROUBLESHOOTING:{RESET}\n- Ensure Unity is open.\n- Check internet.\n- Review errors.")
        if tutor and hasattr(tutor, 'metrics_logger') and tutor.metrics_logger:
             tutor.metrics_logger.log_error(f"Fatal Error during run: {e}")
             tutor.metrics_logger.log_session_end()
             print(f"{BLUE}Attempted to save final metrics.{RESET}")
        elif temp_metrics_logger:
             temp_metrics_logger.log_error(f"Fatal Error before/during tutor run: {e}")
             temp_metrics_logger.log_session_end()
             print(f"{BLUE}Attempted to save final metrics via temp logger.{RESET}")
    finally:
        print(f"\n{CYAN}Tutor session finished.{RESET}")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print(f"\n{YELLOW}Tutor interrupted. Exiting...{RESET}")
    except Exception as e:
        print(f"\n{RED}--- UNEXPECTED TOP-LEVEL ERROR ---{RESET}\n{RED}Error: {e}{RESET}")
        traceback.print_exc()
    finally:
        print(f"\n{MAGENTA}Exiting application.{RESET}")
